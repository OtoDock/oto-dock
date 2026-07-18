"""Tests for the LaTeX equation pipeline: normalization, the validation
guard (bad LaTeX must fail loudly, never corrupt a document), native OMML in
DOCX with LaTeX round-trip on read, math spans in write_pdf (incl. the
dollar-amount and code-block traps), and the PNG path for XLSX.
"""

import asyncio
import sys
from pathlib import Path

import pytest

# Make the parent dir importable as a top-level module
sys.path.insert(0, str(Path(__file__).parent.parent))

pytest.importorskip("latex2mathml")
pytest.importorskip("mathml2omml")

from equations import (
    EquationError,
    latex_to_mathml,
    latex_to_omml_element,
    normalize_latex,
    omml_to_latex,
)


# ---------------------------------------------------------------------------
# Normalization + validation guard
# ---------------------------------------------------------------------------


def test_normalize_strips_delimiters():
    assert normalize_latex(r"$x^2$") == "x^2"
    assert normalize_latex(r"$$x^2$$") == "x^2"
    assert normalize_latex(r"\(x^2\)") == "x^2"
    assert normalize_latex(r"\[x^2\]") == "x^2"


def test_normalize_rewrites_aligned():
    out = normalize_latex(r"\begin{aligned} a &= b \\ c &= d \end{aligned}")
    assert "align*" in out and "aligned" not in out


def test_normalize_brace_wraps_matrix_superscript():
    out = normalize_latex(r"\begin{bmatrix} a & b \\ c & d \end{bmatrix}^{-1}")
    assert out.startswith("{") and "}^{-1}" in out


def test_valid_equations_convert():
    for tex in (
        r"\sum_{i=1}^{n} \frac{x_i^2}{\sigma^2}",
        r"SOC^{min} \leq soc_t \leq SOC^{max} \quad \forall t",
        r"\begin{cases} 1 & x > 0 \\ 0 & \text{else} \end{cases}",
        r"\operatorname{argmin}_x \|Ax - b\|_2^2",
    ):
        assert latex_to_mathml(tex).startswith("<math")


def test_unknown_command_rejected_with_actionable_error():
    with pytest.raises(EquationError, match="undefinedmacro"):
        latex_to_mathml(r"x + \undefinedmacro{q}")


def test_empty_equation_rejected():
    with pytest.raises(EquationError):
        latex_to_mathml("  $$  $$  ")


# ---------------------------------------------------------------------------
# OMML (DOCX native math)
# ---------------------------------------------------------------------------


def test_omml_element_has_word_property_children():
    from lxml import etree

    from equations import OMML_NS

    omath = latex_to_omml_element(r"\frac{a}{b}")
    assert omath.tag == f"{{{OMML_NS}}}oMath"
    frac = omath.find(f".//{{{OMML_NS}}}f")
    assert frac is not None
    assert frac[0].tag == f"{{{OMML_NS}}}fPr"
    etree.tostring(omath)  # serializes cleanly


def test_omml_round_trips_to_latex():
    omath = latex_to_omml_element(r"\sum_{i=1}^{n} \frac{x_i^2}{\sigma^2}")
    back = omml_to_latex(omath)
    assert back and r"\sum" in back and r"\frac" in back


# ---------------------------------------------------------------------------
# DOCX write + read
# ---------------------------------------------------------------------------


@pytest.fixture()
def _word(monkeypatch):
    pytest.importorskip("docx")
    import word as word_mod

    async def _noop_preview(path, filename=None):
        return None

    monkeypatch.setattr(word_mod, "_resolve_path", lambda p, writing=False: p)
    monkeypatch.setattr(word_mod, "_push_preview", _noop_preview)
    return word_mod


def test_docx_equation_roundtrip(tmp_path, _word):
    f = tmp_path / "eq.docx"
    msg = asyncio.run(_word.handle_write_docx({
        "path": str(f),
        "operations": [
            {"type": "add_equation", "latex": r"E = mc^2"},
            {"type": "add_paragraph", "runs": [
                {"text": "Inline "},
                {"equation": r"\frac{p_t^d}{\eta^d}"},
                {"text": " math."},
            ]},
        ],
    }))
    assert "2 operations applied" in msg and "Errors" not in msg
    out = _word.read_docx(str(f))
    assert "$$" in out and "mc^{2}" in out.replace(" ", "")
    assert r"\frac" in out


def test_docx_bad_equation_fails_op_without_leftover_paragraph(tmp_path, _word):
    import docx

    f = tmp_path / "bad.docx"
    msg = asyncio.run(_word.handle_write_docx({
        "path": str(f),
        "operations": [
            {"type": "add_paragraph", "text": "before"},
            {"type": "add_equation", "latex": r"\notarealcmd{x}"},
        ],
    }))
    assert "notarealcmd" in msg
    doc = docx.Document(str(f))
    # The failed op must not leave an empty paragraph behind
    assert [p.text for p in doc.paragraphs] == ["before"]


# ---------------------------------------------------------------------------
# PDF math spans
# ---------------------------------------------------------------------------


def test_pdf_math_span_extraction_rules():
    from pdf import _extract_math_spans

    text = (
        "Cost $100 and $200, range $5-$10.\n"
        "Real math $x^2$ and \\(a+b\\) and $$\\frac{1}{2}$$\n"
        "`code $x$ here` and\n```\nblock $y^2$ money\n```\n"
    )
    replaced, spans = _extract_math_spans(text, is_markdown=True)
    latexes = [s[1] for s in spans]
    # Extraction order follows delimiter precedence: $$ / \[ first, then \(, then $
    assert latexes == [r"\frac{1}{2}", "a+b", "x^2"]
    # Dollar amounts and code content untouched
    assert "$100 and $200, range $5-$10" in replaced
    assert "`code $x$ here`" in replaced
    assert "block $y^2$ money" in replaced


def test_pdf_renders_svg_and_reports_failures(tmp_path, monkeypatch):
    pytest.importorskip("weasyprint")
    pytest.importorskip("ziamath")
    import pdf as pdf_mod

    async def _noop_preview(path, filename=None):
        return None

    monkeypatch.setattr(pdf_mod, "_resolve_path", lambda p, writing=False: p)
    monkeypatch.setattr(pdf_mod, "_push_preview", _noop_preview)

    f = tmp_path / "eq.pdf"
    msg = asyncio.run(pdf_mod.handle_write_pdf({
        "path": str(f),
        "content": "Good: $x^2$\n\nBad: \\(\\nosuchcmd{y}\\)\n",
    }))
    assert f.exists() and f.stat().st_size > 0
    assert "1 equation(s) rendered" in msg
    assert "nosuchcmd" in msg  # failure reported, span left as text


def test_pdf_inline_svg_baseline_style():
    from equations import latex_to_svg

    pytest.importorskip("ziamath")
    svg, yofst = latex_to_svg(r"\frac{a}{b}", display=False)
    assert svg.startswith("<svg")
    assert isinstance(yofst, float)


# ---------------------------------------------------------------------------
# XLSX equation image
# ---------------------------------------------------------------------------


def test_xlsx_equation_anchor_and_comment(tmp_path, monkeypatch):
    pytest.importorskip("openpyxl")
    pytest.importorskip("cairosvg")
    pytest.importorskip("cairocffi")  # needs system libcairo
    import openpyxl

    import excel as excel_mod

    async def _noop_preview(path, filename=None):
        return None

    monkeypatch.setattr(excel_mod, "_resolve_path", lambda p, writing=False: p)
    monkeypatch.setattr(excel_mod, "_push_preview", _noop_preview)

    f = tmp_path / "eq.xlsx"
    msg = asyncio.run(excel_mod.handle_write_xlsx({
        "path": str(f),
        "operations": [
            {"type": "write_cells", "cells": [{"cell": "A1", "value": "Balance:"}]},
            {"type": "add_equation", "cell": "B2",
             "latex": r"soc_t = soc_{t-1} + \eta^c p_t^c \Delta t"},
        ],
    }))
    assert "Errors" not in msg
    wb = openpyxl.load_workbook(str(f))
    ws = wb.active
    assert ws["B2"].comment and "LaTeX:" in ws["B2"].comment.text
    assert len(ws._images) == 1
