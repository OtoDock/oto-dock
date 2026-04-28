"""Word (DOCX) read and write handlers for the file-tools MCP.

Enhanced reader with document properties, sections, headers/footers, hyperlinks.
22 write operations including headers/footers, hyperlinks, TOC, page numbers,
search & replace, table manipulation, and mixed-formatting paragraphs.
"""

from pathlib import Path

from shared import _normalize_operations, _op_type, _push_preview, _resolve_path, _to_agents_relative, logger


# ---------------------------------------------------------------------------
# Read
# ---------------------------------------------------------------------------


def read_docx(path: str) -> str:
    from docx import Document
    from docx.shared import Emu

    doc = Document(path)
    result = [f"**DOCX**: {Path(path).name}"]
    result.append("")

    # Document properties
    try:
        props = doc.core_properties
        meta = []
        if props.title:
            meta.append(f"Title: {props.title}")
        if props.author:
            meta.append(f"Author: {props.author}")
        if props.subject:
            meta.append(f"Subject: {props.subject}")
        if meta:
            result.append("**Properties**: " + " | ".join(meta))
    except Exception:
        pass

    # Section info
    try:
        for si, section in enumerate(doc.sections):
            orient = "landscape" if section.orientation == 1 else "portrait"
            w = round(section.page_width / Emu(914400), 1) if section.page_width else "?"
            h = round(section.page_height / Emu(914400), 1) if section.page_height else "?"
            has_header = bool(section.header and section.header.paragraphs and
                             any(p.text.strip() for p in section.header.paragraphs))
            has_footer = bool(section.footer and section.footer.paragraphs and
                             any(p.text.strip() for p in section.footer.paragraphs))
            hf = []
            if has_header:
                hf.append("header")
            if has_footer:
                hf.append("footer")
            hf_str = f", {'+'.join(hf)}" if hf else ""
            if len(doc.sections) > 1:
                result.append(f"**Section {si+1}**: {w}x{h}in {orient}{hf_str}")
            else:
                result.append(f"**Page**: {w}x{h}in {orient}{hf_str}")
    except Exception:
        pass

    # Headers/footers content
    try:
        for si, section in enumerate(doc.sections):
            prefix = f"Section {si+1} " if len(doc.sections) > 1 else ""
            if section.header and section.header.paragraphs:
                hdr_text = " | ".join(
                    p.text.strip() for p in section.header.paragraphs if p.text.strip()
                )
                if hdr_text:
                    result.append(f"**{prefix}Header**: {hdr_text}")
            if section.footer and section.footer.paragraphs:
                ftr_text = " | ".join(
                    p.text.strip() for p in section.footer.paragraphs if p.text.strip()
                )
                if ftr_text:
                    result.append(f"**{prefix}Footer**: {ftr_text}")
    except Exception:
        pass

    # Stats
    img_count = 0
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_count += 1
    except Exception:
        pass
    stats = [f"{len(doc.paragraphs)} paragraphs", f"{len(doc.tables)} tables"]
    if img_count:
        stats.append(f"{img_count} images")
    result.append(f"**Content**: {', '.join(stats)}")
    result.append("")

    # Heading outline
    headings = []
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        if "Heading" in style:
            level = style.replace("Heading", "").strip() or "1"
            headings.append(f"{'  ' * (int(level) - 1)}{level}. {para.text.strip()}")
    if headings:
        result.append("**Outline**:")
        result.extend(headings)
        result.append("")

    # Body content
    result.append("**Body**:")
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        prefix = ""
        if "Heading" in style:
            level = style.replace("Heading", "").strip() or "1"
            prefix = "#" * int(level) + " "
        text = para.text.strip()
        if text:
            result.append(f"{prefix}{text}")

    # Tables
    for i, table in enumerate(doc.tables):
        result.append(f"\n**Table {i+1}** ({len(table.rows)} rows x {len(table.columns)} cols):")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            result.append("| " + " | ".join(cells) + " |")

    return "\n".join(result)


# ---------------------------------------------------------------------------
# Write — helpers
# ---------------------------------------------------------------------------


def _apply_run_format(run, fmt: dict):
    """Apply formatting to a run from a dict."""
    from docx.shared import Pt, RGBColor

    if fmt.get("bold"):
        run.bold = True
    if fmt.get("italic"):
        run.italic = True
    if fmt.get("underline"):
        run.underline = True
    if fmt.get("strikethrough"):
        run.font.strike = True
    if fmt.get("font_name"):
        run.font.name = fmt["font_name"]
    if fmt.get("font_size"):
        run.font.size = Pt(int(fmt["font_size"]))
    color = fmt.get("color")
    if color:
        c = str(color).lstrip("#")
        run.font.color.rgb = RGBColor(int(c[:2], 16), int(c[2:4], 16), int(c[4:6], 16))


def _apply_paragraph_spacing(p, op: dict) -> list[str]:
    """Apply paragraph-level spacing and indentation from op dict.

    UNITS (a common source of confusion):
      * ``space_before`` / ``space_after`` — POINTS (pt). NOT twips/EMUs.
        e.g. 9 ≈ a blank half-line; 72 = one inch.
      * ``line_spacing`` — a MULTIPLE of the line height (1.0 = single,
        1.5, 2.0 = double). NOT points.
      * ``first_line_indent`` / ``left_indent`` — INCHES.

    Returns a list of human-readable warnings (e.g. an implausibly large
    spacing value) for ``write_docx`` to surface back to the caller.
    """
    from docx.shared import Pt, Inches

    warnings: list[str] = []
    pf = p.paragraph_format
    # A US-Letter page is 792pt tall. A space_before/after near or above that
    # almost always means the unit was confused (e.g. twips passed as points),
    # which blows each paragraph onto its own page — warn loudly.
    _MAX_REASONABLE_PT = 500
    if op.get("space_before") is not None:
        v = float(op["space_before"])
        pf.space_before = Pt(v)
        if v > _MAX_REASONABLE_PT:
            warnings.append(
                f"space_before={v:g}pt (~{v / 72:.1f} in) is larger than a page — "
                "space_before/space_after are in POINTS, not twips; did you mean a "
                "smaller value (e.g. 6–18)?"
            )
    if op.get("space_after") is not None:
        v = float(op["space_after"])
        pf.space_after = Pt(v)
        if v > _MAX_REASONABLE_PT:
            warnings.append(
                f"space_after={v:g}pt (~{v / 72:.1f} in) is larger than a page — "
                "space_before/space_after are in POINTS, not twips; did you mean a "
                "smaller value (e.g. 6–18)?"
            )
    if op.get("line_spacing") is not None:
        ls = float(op["line_spacing"])
        pf.line_spacing = ls
        if ls > 10:
            warnings.append(
                f"line_spacing={ls:g} is a MULTIPLE of the line height (1.0=single, "
                "2.0=double), not points — a value this large is almost certainly a "
                "mistake."
            )
    if op.get("first_line_indent") is not None:
        pf.first_line_indent = Inches(float(op["first_line_indent"]))
    if op.get("left_indent") is not None:
        pf.left_indent = Inches(float(op["left_indent"]))
    if op.get("keep_together") is not None:
        pf.keep_together = op["keep_together"]
    if op.get("keep_with_next") is not None:
        pf.keep_with_next = op["keep_with_next"]
    if op.get("widow_control") is not None:
        pf.widow_control = op["widow_control"]
    return warnings


def _add_hyperlink_to_paragraph(paragraph, url: str, text: str,
                                 color: str = "0563C1", underline: bool = True):
    """Add a hyperlink to a paragraph using OxmlElement."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    if color:
        c_elem = OxmlElement("w:color")
        c_elem.set(qn("w:val"), color.lstrip("#"))
        rPr.append(c_elem)
    if underline:
        u_elem = OxmlElement("w:u")
        u_elem.set(qn("w:val"), "single")
        rPr.append(u_elem)

    # Use hyperlink style
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    new_run.append(rPr)
    t_elem = OxmlElement("w:t")
    t_elem.set(qn("xml:space"), "preserve")
    t_elem.text = text
    new_run.append(t_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_field_code(run, field_code: str):
    """Insert a Word field code (PAGE, NUMPAGES, etc.) into a run."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    # Placeholder text
    t_elem = OxmlElement("w:t")
    t_elem.text = "#"
    run._r.append(t_elem)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _search_replace_in_paragraph(paragraph, find: str, replace: str,
                                  match_case: bool = True) -> int:
    """Replace text in a paragraph, handling text split across runs."""
    full_text = paragraph.text
    check_text = full_text if match_case else full_text.lower()
    check_find = find if match_case else find.lower()

    if check_find not in check_text:
        return 0

    # Try simple case first: text in a single run
    count = 0
    for run in paragraph.runs:
        run_text = run.text if match_case else run.text.lower()
        if check_find in (run.text if match_case else run.text.lower()):
            if match_case:
                run.text = run.text.replace(find, replace)
            else:
                # Case-insensitive: use position-based replacement
                import re
                run.text = re.sub(re.escape(find), replace, run.text, flags=re.IGNORECASE)
            count += run.text.count(replace) if count == 0 else 0
            count = max(count, 1)
            return count

    # Text spans multiple runs: rebuild
    # Collect all run texts and their positions
    runs = paragraph.runs
    if not runs:
        return 0

    combined = "".join(r.text for r in runs)
    combined_check = combined if match_case else combined.lower()

    if check_find not in combined_check:
        return 0

    # Simple approach: replace in combined text, then redistribute
    if match_case:
        new_text = combined.replace(find, replace)
    else:
        import re
        new_text = re.sub(re.escape(find), replace, combined, flags=re.IGNORECASE)

    count = (len(combined) - len(combined.replace(check_find, ""))) // len(check_find)

    # Put all text in first run, clear others
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""

    return max(count, 1)


# ---------------------------------------------------------------------------
# Write — main handler
# ---------------------------------------------------------------------------


async def handle_write_docx(args: dict) -> str:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    path = _resolve_path(args["path"], writing=True)
    ops = _normalize_operations(args.get("operations"))
    create_new = args.get("create_new", False)

    if Path(path).exists() and not create_new:
        doc = Document(path)
    else:
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        doc = Document()

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    errors = []

    for idx, op in enumerate(ops):
        ot = _op_type(op)
        try:
            # =============================================================
            # HEADINGS & PARAGRAPHS
            # =============================================================

            if ot == "add_heading":
                doc.add_heading(op["text"], level=int(op.get("level", 1)))

            elif ot == "add_paragraph":
                p = doc.add_paragraph()
                # Support mixed-formatting runs array
                runs_data = op.get("runs")
                if runs_data and isinstance(runs_data, list):
                    for rd in runs_data:
                        run = p.add_run(rd.get("text", ""))
                        _apply_run_format(run, rd)
                else:
                    run = p.add_run(op.get("text", ""))
                    _apply_run_format(run, op)

                if op.get("alignment") and op["alignment"] in align_map:
                    p.alignment = align_map[op["alignment"]]
                for _w in _apply_paragraph_spacing(p, op):
                    errors.append(f"Op #{idx} {ot}: {_w}")

            elif ot == "insert_paragraph":
                index = int(op.get("index", 0))
                paras = doc.paragraphs
                if index < 0 or index > len(paras):
                    errors.append(f"Op #{idx} insert_paragraph: index {index} out of range")
                    continue

                # Create new paragraph element
                new_p = OxmlElement("w:p")
                if index < len(paras):
                    paras[index]._element.addprevious(new_p)
                else:
                    doc.element.body.append(new_p)

                # Wrap it and add content
                from docx.text.paragraph import Paragraph
                paragraph = Paragraph(new_p, doc)
                run = paragraph.add_run(op.get("text", ""))
                _apply_run_format(run, op)
                if op.get("alignment") and op["alignment"] in align_map:
                    paragraph.alignment = align_map[op["alignment"]]

            elif ot == "delete_paragraph":
                index = int(op.get("index", 0))
                paras = doc.paragraphs
                if 0 <= index < len(paras):
                    elem = paras[index]._element
                    elem.getparent().remove(elem)
                else:
                    errors.append(f"Op #{idx} delete_paragraph: index {index} out of range")

            elif ot == "set_paragraph_format":
                index = int(op.get("index", 0))
                paras = doc.paragraphs
                if 0 <= index < len(paras):
                    p = paras[index]
                    if op.get("alignment") and op["alignment"] in align_map:
                        p.alignment = align_map[op["alignment"]]
                    for _w in _apply_paragraph_spacing(p, op):
                        errors.append(f"Op #{idx} {ot}: {_w}")
                else:
                    errors.append(f"Op #{idx} set_paragraph_format: index {index} out of range")

            # =============================================================
            # LISTS
            # =============================================================

            elif ot == "add_list":
                items = op.get("items", [])
                ordered = op.get("ordered", False)
                default_level = int(op.get("level", 1))
                base_style = "List Number" if ordered else "List Bullet"
                for item in items:
                    # Items can be plain strings or dicts with text+level
                    if isinstance(item, dict):
                        text = item.get("text", "")
                        lvl = int(item.get("level", default_level))
                    else:
                        text = str(item)
                        lvl = default_level
                    style = base_style if lvl <= 1 else f"{base_style} {lvl}"
                    try:
                        doc.add_paragraph(text, style=style)
                    except Exception:
                        doc.add_paragraph(text, style=base_style)

            # =============================================================
            # TABLES
            # =============================================================

            elif ot == "add_table":
                headers = op.get("headers", [])
                rows = op.get("rows", [])
                style = op.get("style", "Table Grid")
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                try:
                    table.style = style
                except Exception:
                    pass
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = str(h)
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        if ci < len(headers):
                            table.rows[ri + 1].cells[ci].text = str(val)

            elif ot == "merge_table_cells":
                ti = int(op.get("table_index", 0))
                if ti < 0 or ti >= len(doc.tables):
                    errors.append(f"Op #{idx} merge_table_cells: table_index {ti} out of range")
                    continue
                table = doc.tables[ti]
                sr, sc = int(op["start_row"]), int(op["start_col"])
                er, ec = int(op["end_row"]), int(op["end_col"])
                table.cell(sr, sc).merge(table.cell(er, ec))

            elif ot == "add_table_row":
                ti = int(op.get("table_index", -1))
                if ti < 0:
                    ti = len(doc.tables) - 1
                if ti < 0 or ti >= len(doc.tables):
                    errors.append(f"Op #{idx} add_table_row: table_index {ti} out of range")
                    continue
                table = doc.tables[ti]
                new_row = table.add_row()
                values = op.get("values", [])
                for ci, val in enumerate(values):
                    if ci < len(new_row.cells):
                        new_row.cells[ci].text = str(val)

            elif ot == "delete_table_row":
                ti = int(op.get("table_index", -1))
                if ti < 0:
                    ti = len(doc.tables) - 1
                if ti < 0 or ti >= len(doc.tables):
                    errors.append(f"Op #{idx} delete_table_row: table_index out of range")
                    continue
                table = doc.tables[ti]
                row_idx = int(op.get("row_index", -1))
                if row_idx < 0:
                    row_idx = len(table.rows) - 1
                if 0 <= row_idx < len(table.rows):
                    tr = table.rows[row_idx]._tr
                    tr.getparent().remove(tr)
                else:
                    errors.append(f"Op #{idx} delete_table_row: row_index {row_idx} out of range")

            elif ot == "set_table_column_widths":
                ti = int(op.get("table_index", -1))
                if ti < 0:
                    ti = len(doc.tables) - 1
                if ti < 0 or ti >= len(doc.tables):
                    errors.append(f"Op #{idx} set_table_column_widths: table_index out of range")
                    continue
                table = doc.tables[ti]
                widths = op.get("widths", [])
                for ci, w in enumerate(widths):
                    if ci < len(table.columns):
                        table.columns[ci].width = Inches(float(w))

            elif ot == "style_table":
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls

                ti = int(op.get("table_index", -1))
                if ti < 0:
                    ti = len(doc.tables) - 1
                if ti < 0 or ti >= len(doc.tables):
                    errors.append(f"Op #{idx} style_table: table_index {ti} out of range")
                    continue
                table = doc.tables[ti]

                header = op.get("header", {})
                if header and table.rows:
                    hdr_bg = str(header.get("fill") or header.get("background") or "").lstrip("#")
                    hdr_fc = str(header.get("font_color") or header.get("color") or "").lstrip("#")
                    hdr_bold = header.get("bold", True)
                    for cell in table.rows[0].cells:
                        if hdr_bg:
                            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hdr_bg}"/>')
                            cell._tc.get_or_add_tcPr().append(shading)
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if hdr_bold:
                                    run.bold = True
                                if hdr_fc:
                                    run.font.color.rgb = RGBColor(
                                        int(hdr_fc[:2], 16), int(hdr_fc[2:4], 16), int(hdr_fc[4:6], 16)
                                    )

                alt_colors = op.get("alternating_colors") or op.get("stripe_colors")
                if alt_colors and isinstance(alt_colors, list) and len(alt_colors) >= 2:
                    for ri, row in enumerate(table.rows):
                        if ri == 0 and header:
                            continue
                        color = str(alt_colors[ri % len(alt_colors)]).lstrip("#")
                        if color:
                            for cell in row.cells:
                                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
                                cell._tc.get_or_add_tcPr().append(shading)

                row_fill = op.get("row_fill")
                if row_fill and isinstance(row_fill, dict):
                    for row_str, color in row_fill.items():
                        ri = int(row_str)
                        if 0 <= ri < len(table.rows):
                            c = str(color).lstrip("#")
                            for cell in table.rows[ri].cells:
                                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{c}"/>')
                                cell._tc.get_or_add_tcPr().append(shading)

                border_val = op.get("border")
                if border_val:
                    bcolor = "000000"
                    bsize = "4"
                    if isinstance(border_val, dict):
                        bcolor = str(border_val.get("color", "000000")).lstrip("#")
                        bsize = str(border_val.get("size", "4"))
                    tbl = table._tbl
                    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
                    borders_xml = (
                        f'<w:tblBorders {nsdecls("w")}>'
                        f'<w:top w:val="single" w:sz="{bsize}" w:color="{bcolor}"/>'
                        f'<w:left w:val="single" w:sz="{bsize}" w:color="{bcolor}"/>'
                        f'<w:bottom w:val="single" w:sz="{bsize}" w:color="{bcolor}"/>'
                        f'<w:right w:val="single" w:sz="{bsize}" w:color="{bcolor}"/>'
                        f'<w:insideH w:val="single" w:sz="{bsize}" w:color="{bcolor}"/>'
                        f'<w:insideV w:val="single" w:sz="{bsize}" w:color="{bcolor}"/>'
                        f"</w:tblBorders>"
                    )
                    tblPr.append(parse_xml(borders_xml))

            # =============================================================
            # IMAGES
            # =============================================================

            elif ot == "add_image":
                img_path = _resolve_path(
                    op.get("image_path") or op.get("path") or op.get("image", "")
                )
                width = Inches(float(op.get("width_inches", 4.0)))
                height = Inches(float(op["height_inches"])) if op.get("height_inches") else None
                pic = doc.add_picture(img_path, width=width, height=height)
                # Center image if requested
                alignment = op.get("alignment")
                if alignment and alignment in align_map:
                    pic_paragraph = doc.paragraphs[-1]
                    pic_paragraph.alignment = align_map[alignment]

            # =============================================================
            # PAGE STRUCTURE
            # =============================================================

            elif ot == "add_page_break":
                doc.add_page_break()

            elif ot == "add_section_break":
                from docx.enum.section import WD_ORIENT

                new_section = doc.add_section()
                if op.get("orientation") == "landscape":
                    new_section.orientation = WD_ORIENT.LANDSCAPE
                    new_section.page_width, new_section.page_height = (
                        new_section.page_height, new_section.page_width
                    )
                elif op.get("orientation") == "portrait":
                    new_section.orientation = WD_ORIENT.PORTRAIT
                for margin in ("top", "bottom", "left", "right"):
                    key = f"margin_{margin}"
                    if key in op:
                        setattr(new_section, f"{margin}_margin", Inches(float(op[key])))

            elif ot == "set_page_setup":
                section = doc.sections[-1]
                if op.get("orientation") == "landscape":
                    section.orientation = 1
                    section.page_width, section.page_height = (
                        section.page_height, section.page_width
                    )
                for margin in ("top", "bottom", "left", "right"):
                    key = f"margin_{margin}"
                    if key in op:
                        setattr(section, f"{margin}_margin", Inches(float(op[key])))

            # =============================================================
            # HEADERS & FOOTERS
            # =============================================================

            elif ot == "add_header":
                si = int(op.get("section_index", -1))
                section = doc.sections[si]
                header = section.header
                header.is_linked_to_previous = False

                # Clear existing paragraphs
                for p in header.paragraphs:
                    if not p.text.strip():
                        continue
                p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                p.clear()

                runs_data = op.get("runs")
                if runs_data and isinstance(runs_data, list):
                    for rd in runs_data:
                        run = p.add_run(rd.get("text", ""))
                        _apply_run_format(run, rd)
                else:
                    run = p.add_run(op.get("text", ""))
                    _apply_run_format(run, op)

                if op.get("alignment") and op["alignment"] in align_map:
                    p.alignment = align_map[op["alignment"]]

            elif ot == "add_footer":
                si = int(op.get("section_index", -1))
                section = doc.sections[si]
                footer = section.footer
                footer.is_linked_to_previous = False

                p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                p.clear()

                runs_data = op.get("runs")
                if runs_data and isinstance(runs_data, list):
                    for rd in runs_data:
                        run = p.add_run(rd.get("text", ""))
                        _apply_run_format(run, rd)
                else:
                    run = p.add_run(op.get("text", ""))
                    _apply_run_format(run, op)

                if op.get("alignment") and op["alignment"] in align_map:
                    p.alignment = align_map[op["alignment"]]

            elif ot == "add_page_number":
                si = int(op.get("section_index", -1))
                section = doc.sections[si]
                footer = section.footer
                footer.is_linked_to_previous = False

                p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                # Don't clear — allow combining with footer text
                if p.text.strip():
                    p = footer.add_paragraph()

                position = op.get("position", "center")
                if position in align_map:
                    p.alignment = align_map[position]

                fmt = op.get("format", "{page}")
                font_size = int(op.get("font_size", 10))

                # Parse format string and insert field codes
                parts = fmt.replace("{page}", "\x00PAGE\x00").replace("{total}", "\x00NUMPAGES\x00").split("\x00")
                for part in parts:
                    if part == "PAGE":
                        run = p.add_run()
                        if font_size:
                            run.font.size = Pt(font_size)
                        _add_field_code(run, "PAGE")
                    elif part == "NUMPAGES":
                        run = p.add_run()
                        if font_size:
                            run.font.size = Pt(font_size)
                        _add_field_code(run, "NUMPAGES")
                    elif part:
                        run = p.add_run(part)
                        if font_size:
                            run.font.size = Pt(font_size)

            # =============================================================
            # HYPERLINKS
            # =============================================================

            elif ot == "add_hyperlink":
                url = op.get("url", "")
                text = op.get("text", url)
                color = op.get("color", "0563C1")
                underline = op.get("underline", True)
                pi = op.get("paragraph_index")

                if pi is not None:
                    paras = doc.paragraphs
                    pidx = int(pi)
                    if 0 <= pidx < len(paras):
                        paragraph = paras[pidx]
                    else:
                        paragraph = doc.add_paragraph()
                else:
                    paragraph = doc.add_paragraph()

                _add_hyperlink_to_paragraph(paragraph, url, text, color, underline)

            # =============================================================
            # TABLE OF CONTENTS
            # =============================================================

            elif ot == "add_table_of_contents":
                title = op.get("title", "Table of Contents")
                levels = int(op.get("levels", 3))

                if title:
                    doc.add_heading(title, level=1)

                p = doc.add_paragraph()
                run = p.add_run()
                # Begin field
                fld_begin = OxmlElement("w:fldChar")
                fld_begin.set(qn("w:fldCharType"), "begin")
                run._r.append(fld_begin)

                instr = OxmlElement("w:instrText")
                instr.set(qn("xml:space"), "preserve")
                instr.text = f' TOC \\o "1-{levels}" \\h \\z \\u '
                run._r.append(instr)

                fld_sep = OxmlElement("w:fldChar")
                fld_sep.set(qn("w:fldCharType"), "separate")
                run._r.append(fld_sep)

                # Placeholder
                run2 = p.add_run("Update table of contents to see entries")
                run2.italic = True
                run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

                fld_end = OxmlElement("w:fldChar")
                fld_end.set(qn("w:fldCharType"), "end")
                run2._r.append(fld_end)

            # =============================================================
            # SEARCH & REPLACE
            # =============================================================

            elif ot == "search_and_replace":
                find_text = op.get("find", "")
                replace_text = op.get("replace", "")
                match_case = op.get("match_case", True)

                if not find_text:
                    errors.append(f"Op #{idx} search_and_replace: 'find' text is empty")
                    continue

                total_replacements = 0
                # Search in body paragraphs
                for p in doc.paragraphs:
                    total_replacements += _search_replace_in_paragraph(
                        p, find_text, replace_text, match_case
                    )
                # Search in table cells
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                total_replacements += _search_replace_in_paragraph(
                                    p, find_text, replace_text, match_case
                                )
                # Search in headers/footers
                for section in doc.sections:
                    if section.header:
                        for p in section.header.paragraphs:
                            total_replacements += _search_replace_in_paragraph(
                                p, find_text, replace_text, match_case
                            )
                    if section.footer:
                        for p in section.footer.paragraphs:
                            total_replacements += _search_replace_in_paragraph(
                                p, find_text, replace_text, match_case
                            )

                if total_replacements == 0:
                    errors.append(f"Op #{idx} search_and_replace: '{find_text}' not found")

            # =============================================================
            # UNKNOWN
            # =============================================================

            else:
                logger.warning(f"write_docx: unknown operation '{ot}', skipping")
                errors.append(f"Op #{idx}: unknown operation '{ot}'")

        except Exception as exc:
            errors.append(f"Op #{idx} {ot}: {exc}")
            logger.warning(f"write_docx op #{idx} '{ot}' failed: {exc}")

    # A trailing 'add_page_break' starts a NEW page after the cursor, leaving a
    # blank final page — almost always unintended. Warn (don't silently drop it,
    # in case a blank page IS wanted).
    if ops and _op_type(ops[-1]) == "add_page_break":
        errors.append(
            "The final operation is 'add_page_break' — it starts a NEW page after "
            "the content and so leaves a trailing BLANK page. Remove it as the last "
            "op unless an intentional blank final page is desired."
        )

    # Save even if some operations failed
    doc.save(path)
    await _push_preview(path)

    msg = f"Document saved: {_to_agents_relative(path)} ({len(ops)} operations applied)"
    if errors:
        msg += f"\n\nWarnings/Errors ({len(errors)}):\n" + "\n".join(f"  - {e}" for e in errors)
    return msg
